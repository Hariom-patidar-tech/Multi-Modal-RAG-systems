import os
from docx import Document
from app.core.logger import logger

class DocxLoader:
    def __init__(self):
        pass

    def load(self, file_path: str) -> str:
        
        if not os.path.exists(file_path):
            logger.error(f"DOCX file not found at path: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        logger.info(f"Extracting content from DOCX: {file_path}")
        try:
            doc = Document(file_path)
            extracted_elements = []

            
            for para in doc.paragraphs:
                if para.text.strip():  
                    extracted_elements.append(para.text.strip())

           
            for table_idx, table in enumerate(doc.tables):
                extracted_elements.append(f"\n--- Table {table_idx + 1} Start ---")
                
                for row in table.rows:
                    
                    row_text = [cell.text.strip() for cell in row.cells]
                    row_string = " | ".join(row_text)
                    if row_string.strip():
                        extracted_elements.append(row_string)
                        
                extracted_elements.append(f"--- Table {table_idx + 1} End ---\n")

            final_text = "\n".join(extracted_elements)
            logger.info(f"Successfully extracted DOCX content. Total length: {len(final_text)} characters.")
            return final_text

        except Exception as e:
            logger.error(f"Error while parsing DOCX file {file_path}: {str(e)}")
            raise Exception(f"DOCX parsing failed: {str(e)}")